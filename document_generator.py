import os
import re
import copy
from datetime import datetime
from docx import Document
from docx.table import Table

def _shorten_address(address):
    """Strip redundant parts from a Nominatim address for display in documents."""
    if not address:
        return address
    addr = re.sub(r',\s*(Italy|Italia)\s*$', '', address, flags=re.IGNORECASE).strip()
    addr = re.sub(r',\s*Trentino[^,]*', '', addr, flags=re.IGNORECASE).strip()
    addr = re.sub(r',\s*(Alto Adige|S[üu]dtirol|South Tyrol|Sud Tirol)[^,]*', '', addr, flags=re.IGNORECASE).strip()
    addr = re.sub(r',\s*\d{5}\b', '', addr).strip()
    addr = re.sub(r',\s*[A-Z]{2}(?=,|$)', '', addr).strip()
    return addr.rstrip(',').strip()


def _replace_text_in_paragraph(paragraph, search_text, replace_text):
    """Safely replace text within a paragraph while trying to preserve formatting."""
    if search_text not in paragraph.text:
        return
        
    for run in paragraph.runs:
        if search_text in run.text:
            run.text = run.text.replace(search_text, replace_text)
            return

    # Fallback if it spans multiple runs (use first run's formatting)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        target_color = first_run.font.color.rgb if (first_run.font and first_run.font.color) else None
        target_size = first_run.font.size if first_run.font else None
        target_bold = first_run.font.bold if first_run.font else None
    else:
        target_color, target_size, target_bold = None, None, None

    new_text = paragraph.text.replace(search_text, replace_text)
    paragraph.clear()
    new_run = paragraph.add_run(new_text)
    if target_color:
        new_run.font.color.rgb = target_color
    if target_size:
        new_run.font.size = target_size
    if target_bold is not None:
        new_run.font.bold = target_bold


def generate_piano_viaggi(template_path, out_docx_path, data):
    """
    data = {
        'event_name': '...',
        'date': '...',
        'destination': '...',
        'start_time': '...',
        'end_time': '...',
        'routes': [...],
        'exclude_autonomia': bool
    }
    """
    doc = Document(template_path)

    date_str = data.get('date', '').strip()
    prep = "dell'" if date_str.startswith(('1 ', '8 ', '11 ')) else "del "
    event_str = f"{data.get('event_name', '')} {prep}{date_str}"

    # 1. Replace Placeholders in Headers
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, 'NOME EVENTO  dell’ DATA', event_str)
        _replace_text_in_paragraph(p, 'SEDE', data.get('destination', ''))
        _replace_text_in_paragraph(p, '8,30', data.get('start_time', '8:30'))
        _replace_text_in_paragraph(p, '16,00', data.get('end_time', '16:00'))
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    _replace_text_in_paragraph(p, 'NOME EVENTO  dell’ DATA', event_str)
                    _replace_text_in_paragraph(p, 'SEDE', data.get('destination', ''))
                    _replace_text_in_paragraph(p, '8,30', data.get('start_time', '8:30'))
                    _replace_text_in_paragraph(p, '16,00', data.get('end_time', '16:00'))

    # 2. Find the marker where to insert dynamic tables
    # and find the Header Table
    header_table_idx = -1
    for i, t in enumerate(doc.tables):
        if len(t.rows) > 0 and t.rows[0].cells[0].text.strip() == "N°":
            header_table_idx = i
            break

    if header_table_idx != -1 and header_table_idx + 1 < len(doc.tables):
        base_tbl_element = copy.deepcopy(doc.tables[header_table_idx + 1]._tbl)
        
        # Remove all mock tables after the header table
        tables_to_remove = []
        for i in range(header_table_idx + 1, len(doc.tables)):
            tables_to_remove.append(doc.tables[i]._element)
        for tbl_elem in tables_to_remove:
            tbl_elem.getparent().remove(tbl_elem)

        # insertion point
        target_p = None
        for p in doc.paragraphs:
            if "Gli Istituti sotto riportati" in p.text:
                target_p = p
                break
        
        if not target_p:
            # Fallback
            target_p = doc.paragraphs[-1]

        # 3. Generate tables for each route, inserting right after the header table
        from docx.oxml.shared import OxmlElement
        last_inserted = doc.tables[header_table_idx]._element

        routes = data.get('routes', [])
        for route_idx, route in enumerate(routes):
            new_tbl_element = copy.deepcopy(base_tbl_element)
            t = Table(new_tbl_element, doc)
            
            # The base table usually has: Data Row 1, Data Row 2, Totals Row
            # Let's dynamically create rows for stops. We'll copy the first row,
            # clear all standard data rows, then append.
            data_row_template = copy.deepcopy(t.rows[0]._tr)
            totals_row_template = copy.deepcopy(t.rows[-1]._tr)
            
            # Remove all rows inside new_tbl_element
            for row_elem in list(new_tbl_element):
                if row_elem.tag.endswith('tr'):
                    new_tbl_element.remove(row_elem)
                    
            # Identify stops (only pickups)
            pickup_stops = [s for s in route.get('outbound', {}).get('stops', []) if s.get('type') == 'pickup']
            total_pax = 0
            
            from docx.shared import Pt
            def _set_cell_val(cell, text, size=10):
                cell.text = text
                if cell.paragraphs:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(size)
                        run.font.name = 'Arial'
            
            first_fin_cell = None
            first_km_cell = None
            first_name_cell = None
            last_name = None
            first_time_cell = None
            last_time = None
            first_ret_time_cell = None
            last_ret_time = None

            for i, stop in enumerate(pickup_stops):
                stop_row_elem = copy.deepcopy(data_row_template)
                new_tbl_element.append(stop_row_elem)
                # re-wrap to modify
                temp_t = Table(new_tbl_element, doc)
                current_row = temp_t.rows[-1]

                if i == 0:
                    _set_cell_val(current_row.cells[0], f"Fin {route.get('vehicle_id', route_idx) + 1}", size=9)
                    _set_cell_val(current_row.cells[1], f"{route.get('outbound', {}).get('distance', 0) / 500:.1f} KM", size=11)
                    first_fin_cell = current_row.cells[0]
                    first_km_cell = current_row.cells[1]
                else:
                    _set_cell_val(current_row.cells[0], "")
                    _set_cell_val(current_row.cells[1], "")
                    first_fin_cell.merge(current_row.cells[0])
                    first_km_cell.merge(current_row.cells[1])

                # Istituto/Name (colonna 2) — merge consecutive identical names
                name = stop.get('name', '')
                if i > 0 and name and name == last_name:
                    _set_cell_val(current_row.cells[2], "")
                    first_name_cell.merge(current_row.cells[2])
                else:
                    _set_cell_val(current_row.cells[2], name)
                    first_name_cell = current_row.cells[2]
                    last_name = name

                # Orario Andata (colonna 3)
                current_time = stop.get('departure_time', '')
                if i > 0 and current_time and current_time == last_time:
                    _set_cell_val(current_row.cells[3], "")
                    first_time_cell.merge(current_row.cells[3])
                else:
                    _set_cell_val(current_row.cells[3], current_time)
                    first_time_cell = current_row.cells[3]
                    last_time = current_time

                _set_cell_val(current_row.cells[4], _shorten_address(stop.get('display_address') or stop.get('address', '')))
                _set_cell_val(current_row.cells[5], str(stop.get('count', '0')))

                # Orario Ritorno (colonna 6)
                current_ret_time = stop.get('return_time', '')
                if i > 0 and current_ret_time and current_ret_time == last_ret_time:
                    _set_cell_val(current_row.cells[6], "")
                    first_ret_time_cell.merge(current_row.cells[6])
                else:
                    _set_cell_val(current_row.cells[6], current_ret_time)
                    first_ret_time_cell = current_row.cells[6]
                    last_ret_time = current_ret_time

                total_pax += stop.get('count', 0)
                
            # Append Totals row
            new_tbl_element.append(totals_row_template)
            temp_t = Table(new_tbl_element, doc)
            totals_row = temp_t.rows[-1]
            _set_cell_val(totals_row.cells[-2], str(total_pax), size=11)
            
            # Insert table immediately after the previous element (no gap)
            last_inserted.addnext(new_tbl_element)
            p_elem = OxmlElement('w:p')
            new_tbl_element.addnext(p_elem)
            last_inserted = p_elem

        # 4. Handle generic exclusions at the end
        if data.get('exclude_autonomia', False):
            # Remove all paragraphs starting from "Gli Istituti"
            found_autonomia = False
            for p in doc.paragraphs:
                if "Gli Istituti sotto riportati" in p.text:
                    found_autonomia = True
                if found_autonomia:
                    # Remove
                    p_element = p._element
                    p_element.getparent().remove(p_element)

    doc.save(out_docx_path)
    return out_docx_path


def generate_richiesta_servizio(template_path, out_docx_path, data):
    # For Richiesta servizio it uses an almost identical structure.
    return generate_piano_viaggi(template_path, out_docx_path, data)
