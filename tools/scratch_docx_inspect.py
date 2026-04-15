import docx
import json

doc = docx.Document('template_documents/Piano_Viaggi_TIPO.docx')
out = []
for p in doc.paragraphs:
    if p.text.strip():
        out.append(p.text.strip())

for i, t in enumerate(doc.tables):
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if p.text.strip():
                    out.append(f"Table {i}: {p.text.strip()}")

print(json.dumps(out, indent=2, ensure_ascii=False))
